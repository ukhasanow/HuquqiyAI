# Yuklangan hujjatdan (PDF/DOCX/rasm) matn ajratish
import base64
import io
import logging

import httpx

from ..config import GEMINI_API_KEY, GEMINI_MODEL, MAX_HUJJAT_BELGILAR

log = logging.getLogger(__name__)

RASM_KENGAYTMALARI = (".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif")
MAX_RASM_HAJMI = 8 * 1024 * 1024
RASM_SOROV_MUDDATI = 120

# Rasm — ko'chirib yozish, tahlil emas. Model matnni "tuzatib" yoki
# qisqartirib qo'ysa, shartnoma tahlili yo'q bandni muhokama qilib qoladi.
_OCR_KORSATMASI = """Bu hujjat surati. Undagi BARCHA matnni AYNAN ko'chirib yoz.

QAT'IY QOIDALAR:
- Hech narsani qisqartirma, umumlashtirma va tushuntirma — faqat ko'chir.
- Band va modda raqamlarini (1.1, 2.3, 4.1) aynan saqla, ular eng muhimi.
- Qatorlar va xatboshilar tuzilishini saqla.
- O'qib bo'lmagan joyni [o'qilmadi] deb belgila, o'ylab topma.
- Izoh, sarlavha yoki xulosa qo'shma — faqat hujjatdagi matn."""


class HujjatXato(Exception):
    """Hujjatni o'qib bo'lmaganda ko'tariladi."""


def rasmdan_oqish_mavjud() -> bool:
    return bool(GEMINI_API_KEY)


def matn_ajrat(fayl_nomi: str, bayt: bytes, mime: str = "") -> str:
    nomi = fayl_nomi.lower()
    if nomi.endswith(".pdf"):
        matn = _pdf_matn(bayt)
    elif nomi.endswith(".docx"):
        matn = _docx_matn(bayt)
    elif nomi.endswith(".txt"):
        matn = bayt.decode("utf-8", errors="replace")
    elif nomi.endswith(RASM_KENGAYTMALARI) or (mime or "").startswith("image/"):
        matn = rasm_matni(bayt, mime or "image/jpeg")
    else:
        raise HujjatXato("Faqat PDF, DOCX, TXT yoki rasm fayllar qabul qilinadi.")

    matn = matn.strip()
    if not matn:
        raise HujjatXato(
            "Hujjatdan matn ajratib bo'lmadi. Skanerlangan hujjat bo'lsa, "
            "uning suratini yuboring — rasmni o'qib beraman."
        )
    return matn[:MAX_HUJJAT_BELGILAR]


def rasm_matni(bayt: bytes, mime: str = "image/jpeg") -> str:
    """Hujjat suratidagi matnni ko'chirib oladi (OCR).

    Bu tahlil emas, faqat ko'chirish: huquqiy xulosani baribir shartnoma yoki
    jarima moduli qonun matni asosida chiqaradi.
    """
    if not bayt:
        raise HujjatXato("Rasm bo'sh.")
    if len(bayt) > MAX_RASM_HAJMI:
        raise HujjatXato("Rasm hajmi 8 MB dan oshmasligi kerak.")
    if not rasmdan_oqish_mavjud():
        raise HujjatXato(
            "Rasmdan o'qish bu serverda sozlanmagan. PDF yoki DOCX yuklang."
        )

    try:
        javob = httpx.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent",
            params={"key": GEMINI_API_KEY},
            json={
                "contents": [{"parts": [
                    {"text": _OCR_KORSATMASI},
                    {"inline_data": {"mime_type": mime,
                                     "data": base64.b64encode(bayt).decode()}},
                ]}],
                # temperature=0: ko'chirish ijodkorlik talab qilmaydi
                "generationConfig": {"temperature": 0, "maxOutputTokens": 8192},
            },
            timeout=RASM_SOROV_MUDDATI,
        )
        javob.raise_for_status()
        nomzodlar = javob.json().get("candidates") or []
        if not nomzodlar:
            return ""
        qismlar = nomzodlar[0].get("content", {}).get("parts") or []
        return "".join(q.get("text", "") for q in qismlar).strip()
    except HujjatXato:
        raise
    except Exception as e:
        log.warning("Rasmdan matn o'qib bo'lmadi: %s", e)
        raise HujjatXato(_rasm_xato_matni(e)) from e


def _rasm_xato_matni(e: Exception) -> str:
    """Xato sababini foydalanuvchiga to'g'ri aytish.

    Limit yoki tarmoq xatosida "suratni yorug'roq oling" deyish noto'g'ri:
    odam aybni o'z rasmidan qidirib, qayta-qayta suratga oladi.
    """
    s = str(e).lower()
    if "429" in s or "rate limit" in s or "quota" in s or "resource_exhausted" in s:
        return ("Rasm o'qish xizmati hozir band (so'rovlar limiti). Bir necha "
                "daqiqadan so'ng urinib ko'ring yoki PDF/DOCX yuklang.")
    if "401" in s or "403" in s or "api key" in s:
        return ("Rasm o'qish xizmati sozlanmagan (kalit noto'g'ri). "
                "PDF yoki DOCX yuklang.")
    if "timeout" in s or "timed out" in s:
        return "Rasm o'qish juda uzoq davom etdi. Kichikroq surat bilan urinib ko'ring."
    return "Rasmdan matnni o'qib bo'lmadi. Suratni yorug'roq va to'liq oling."


def _pdf_matn(bayt: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(bayt))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as e:
        raise HujjatXato(f"PDF faylni o'qib bo'lmadi: {e}")


def _docx_matn(bayt: bytes) -> str:
    import docx

    try:
        d = docx.Document(io.BytesIO(bayt))
        qismlar = [p.text for p in d.paragraphs]
        for jadval in d.tables:
            for qator in jadval.rows:
                qismlar.extend(katak.text for katak in qator.cells)
        return "\n".join(qismlar)
    except Exception as e:
        raise HujjatXato(f"DOCX faylni o'qib bo'lmadi: {e}")
